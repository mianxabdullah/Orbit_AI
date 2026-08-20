import streamlit as st
from pypdf import PdfReader
from docx import Document


def extract_text(uploaded_file):

    if uploaded_file is None:
        return ""

    file_type = uploaded_file.name.split(".")[-1].lower()

    try:
        if file_type == "txt":
            raw = uploaded_file.read()
            try:
                text = raw.decode("utf-8")
            except UnicodeDecodeError:
                text = raw.decode("latin-1")

        elif file_type == "pdf":
            reader = PdfReader(uploaded_file)
            text = "\n".join(
                (page.extract_text() or "") for page in reader.pages
            )

        elif file_type == "docx":
            doc = Document(uploaded_file)

            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)

            text = "\n".join(parts)

        else:
            st.sidebar.warning(f"Unsupported file type: .{file_type}")
            return ""

    except Exception as e:
        st.sidebar.error(f"Couldn't read file: {e}")
        return ""

    if not text.strip():
        st.sidebar.warning(
            "No readable text found in this file (it may be scanned/image-based)."
        )

    return text